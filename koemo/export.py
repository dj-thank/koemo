"""エクスポート: Markdown / PDF(PySide6) / DOCX(python-docx)。"""
import re


def export_markdown(doc_md, path):
    with open(path, "w", encoding="utf-8") as f:
        f.write(doc_md)


def export_pdf(doc_md, path):
    """Markdown を A4 PDF に整形出力（Qt の QTextDocument→QPdfWriter、追加依存なし）。"""
    from PySide6.QtGui import QTextDocument, QPdfWriter, QPageSize, QPageLayout, QFont
    from PySide6.QtCore import QMarginsF
    writer = QPdfWriter(str(path))
    writer.setPageSize(QPageSize(QPageSize.A4))
    writer.setPageMargins(QMarginsF(15, 15, 15, 15), QPageLayout.Unit.Millimeter)
    doc = QTextDocument()
    doc.setDefaultFont(QFont("Yu Gothic UI", 11))   # 日本語フォント
    doc.setMarkdown(doc_md)
    doc.print_(writer)


def _plain(s):
    return s.replace("**", "").replace("`", "").strip()


def export_docx(doc_md, path):
    """Markdown を見出し/箇条書きを保ったまま DOCX 出力。"""
    from docx import Document
    d = Document()
    for raw in doc_md.split("\n"):
        line = raw.rstrip()
        stripped = line.lstrip()
        if not line:
            continue
        if line.startswith("### "):
            d.add_heading(_plain(line[4:]), level=3)
        elif line.startswith("## "):
            d.add_heading(_plain(line[3:]), level=2)
        elif line.startswith("# "):
            d.add_heading(_plain(line[2:]), level=1)
        elif stripped.startswith(("- ", "* ")):
            d.add_paragraph(_plain(stripped[2:]), style="List Bullet")
        elif re.match(r"\d+\.\s", stripped):
            d.add_paragraph(_plain(re.sub(r"^\d+\.\s", "", stripped)), style="List Number")
        elif set(line) <= set("-—–= "):   # 水平線
            continue
        else:
            d.add_paragraph(_plain(line))
    d.save(str(path))
